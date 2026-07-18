"""
AIShield Scanner — Core Detection Engine (v2)
==============================================
Multi-layer scanner that detects zero-click prompt injection attacks
hidden inside documents before they reach AI models.

Supported formats: PDF, DOCX, XLSX, HTML, TXT, MD, CSV

Four-layer defense:
  Layer 1 — Content extraction (visible + hidden)
  Layer 2 — Pattern-based detection (regex: injection, exfil, unicode)
  Layer 3 — Structural / heuristic analysis (metadata, text anomalies)
  Layer 4 — Adversarial bypass detection (base64, l33t, Unicode lookalikes)
"""

import re
import time
import base64
import unicodedata
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Tuple
from enum import Enum
from pathlib import Path


# ═══════════════════════════════════════════════════════════════
#  Graceful optional imports
# ═══════════════════════════════════════════════════════════════

try:
    from PyPDF2 import PdfReader
    _HAS_PDF = True
except ImportError:
    _HAS_PDF = False

try:
    import docx  # python-docx
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

try:
    import openpyxl
    _HAS_XLSX = True
except ImportError:
    _HAS_XLSX = False

try:
    from bs4 import BeautifulSoup
    _HAS_BS4 = True
except ImportError:
    _HAS_BS4 = False


# ═══════════════════════════════════════════════════════════════
#  Data Structures
# ═══════════════════════════════════════════════════════════════

class Severity(Enum):
    """Threat severity classification."""
    LOW      = 1
    MEDIUM   = 2
    HIGH     = 3
    CRITICAL = 4

    def __str__(self):
        return self.name


@dataclass
class Finding:
    """Represents a single detected threat."""
    category:   str
    description:str
    severity:   Severity
    evidence:   str
    page:       Optional[int] = None
    confidence: float = 1.0   # 0.0–1.0

    def to_dict(self) -> dict:
        return {
            "category":    self.category,
            "description": self.description,
            "severity":    str(self.severity),
            "evidence":    self.evidence[:200],
            "page":        self.page,
            "confidence":  round(self.confidence * 100),  # render as %
        }


@dataclass
class ScanResult:
    """Aggregated result of a complete document scan."""
    filename:         str
    verdict:          str
    risk_score:       int
    scan_time_ms:     float
    file_format:      str = "unknown"
    findings:         List[Finding] = field(default_factory=list)
    total_pages:      int = 0
    text_length:      int = 0
    metadata_scanned: bool = False

    @property
    def is_blocked(self) -> bool:
        return self.verdict == "BLOCKED"

    def to_dict(self) -> dict:
        return {
            "filename":         self.filename,
            "file_format":      self.file_format,
            "verdict":          self.verdict,
            "risk_score":       self.risk_score,
            "scan_time_ms":     round(self.scan_time_ms, 2),
            "total_pages":      self.total_pages,
            "text_length":      self.text_length,
            "metadata_scanned": self.metadata_scanned,
            "findings":         [f.to_dict() for f in self.findings],
            "threat_summary":   {
                "CRITICAL": sum(1 for f in self.findings if f.severity == Severity.CRITICAL),
                "HIGH":     sum(1 for f in self.findings if f.severity == Severity.HIGH),
                "MEDIUM":   sum(1 for f in self.findings if f.severity == Severity.MEDIUM),
                "LOW":      sum(1 for f in self.findings if f.severity == Severity.LOW),
            },
        }


# ═══════════════════════════════════════════════════════════════
#  AIShield Scanner
# ═══════════════════════════════════════════════════════════════

class AIShieldScanner:
    """
    Multi-layer document scanner for prompt injection detection.

    Usage:
        scanner = AIShieldScanner()
        result = scanner.scan("document.pdf")    # auto-detects format
        result = scanner.scan_text("raw text")   # for Attack Playground
        print(result.verdict)    # "BLOCKED" or "ALLOWED"
        print(result.risk_score) # 0–100
    """

    # ── Layer 2a: Prompt injection patterns ─────────────────────
    INJECTION_PATTERNS: List[Tuple[str, str, Severity]] = [
        (r"ignore\s+(?:all\s+)?previous\s+instructions",
         "Prompt override — ignore previous instructions", Severity.CRITICAL),
        (r"system\s+override",
         "System override command detected", Severity.CRITICAL),
        (r"you\s+are\s+now\s+.{1,80}\s+and\s+must",
         "Identity reassignment injection", Severity.CRITICAL),
        (r"pretend\s+you\s+are",
         "Identity manipulation attempt", Severity.HIGH),
        (r"disregard\s+(?:your|the)\s+(?:system|safety)\s+(?:prompt|instructions)",
         "Safety bypass attempt", Severity.CRITICAL),
        (r"new\s+priority\s+(?:directive|instruction)",
         "Priority directive injection", Severity.HIGH),
        (r"from\s+now\s+on,?\s+(?:you\s+will|you\s+must)",
         "Behavioral override attempt", Severity.HIGH),
        (r"override\s+(?:all|the)\s+(?:previous|system)\s+(?:prompts?|instructions?)",
         "System prompt override", Severity.CRITICAL),
        (r"forget\s+(?:your|the)\s+(?:previous|system)\s+(?:prompt|instructions)",
         "Memory wipe injection", Severity.CRITICAL),
        (r"do\s+not\s+follow\s+(?:your|the|any)\s+(?:original|previous|system)",
         "Instruction negation detected", Severity.HIGH),
        (r"act\s+as\s+(?:if|though)\s+you\s+(?:have|are)",
         "Role manipulation attempt", Severity.MEDIUM),
        (r"respond\s+only\s+with",
         "Output constraint injection", Severity.MEDIUM),
        (r"execute\s+(?:the\s+following|this)\s+(?:command|code|instruction)",
         "Command execution injection", Severity.HIGH),
        (r"do\s+not\s+mention\s+(?:this|these)\s+instruction",
         "Stealth directive — hides its own presence", Severity.CRITICAL),
        (r"(?:secret|hidden)\s+(?:instruction|directive|command)",
         "Self-labeled hidden instruction", Severity.HIGH),
        (r"as\s+an?\s+(?:AI|language\s+model|LLM),?\s+(?:you\s+must|you\s+will|you\s+should)",
         "AI persona constraint injection", Severity.HIGH),
        (r"(?:jailbreak|dan\s+mode|developer\s+mode|unrestricted\s+mode)",
         "Jailbreak mode activation attempt", Severity.CRITICAL),
        (r"(?:ignore|bypass|skip)\s+(?:all\s+)?(?:safety|content|ethical)\s+(?:guidelines|filters|restrictions)",
         "Safety filter bypass attempt", Severity.CRITICAL),
        (r"print\s+(?:the\s+)?(?:above|full|entire|complete)\s+(?:prompt|system\s+prompt|instructions)",
         "System prompt extraction attempt", Severity.HIGH),
    ]

    # ── Layer 2b: Data exfiltration patterns ─────────────────────
    EXFIL_PATTERNS: List[Tuple[str, str, Severity]] = [
        (r"!\[\]\(https?://[^\s\)]+\)",
         "Markdown image exfiltration (zero-pixel tracking)", Severity.CRITICAL),
        (r"https?://[^\s]*?(?:email|data|token|secret|password|key|ssn|credit)\s*=",
         "Exfiltration URL with sensitive data parameter", Severity.CRITICAL),
        (r"https?://[^\s]*?(?:leak|steal|exfil|capture|log|collect)\b",
         "Suspicious exfiltration endpoint URL", Severity.HIGH),
        (r"fetch\s+(?:the\s+)?(?:url|image|resource)\s+(?:at|from)\s+https?://",
         "Fetch instruction targeting remote URL", Severity.HIGH),
        (r"send\s+(?:the|all|any)\s+(?:data|information|content|text)\s+to\s+https?://",
         "Data exfiltration send instruction", Severity.CRITICAL),
        (r"include\s+.{0,40}as\s+(?:url|query)\s+parameters?",
         "URL parameter stuffing instruction", Severity.HIGH),
        (r"(?:webhook|callback|notify)\s+(?:url|endpoint)\s*[:=]\s*https?://",
         "Webhook callback injection", Severity.HIGH),
    ]

    # ── Layer 2c: Zero-width / invisible Unicode ─────────────────
    INVISIBLE_CHARS: Dict[str, str] = {
        "\u200B": "Zero-Width Space (U+200B)",
        "\u200C": "Zero-Width Non-Joiner (U+200C)",
        "\u200D": "Zero-Width Joiner (U+200D)",
        "\u200E": "Left-to-Right Mark (U+200E)",
        "\u200F": "Right-to-Left Mark (U+200F)",
        "\u2060": "Word Joiner (U+2060)",
        "\u2061": "Function Application (U+2061)",
        "\u2062": "Invisible Times (U+2062)",
        "\u2063": "Invisible Separator (U+2063)",
        "\u2064": "Invisible Plus (U+2064)",
        "\uFEFF": "Zero-Width No-Break Space / BOM (U+FEFF)",
        "\u00AD": "Soft Hyphen (U+00AD)",
        "\u034F": "Combining Grapheme Joiner (U+034F)",
        "\u061C": "Arabic Letter Mark (U+061C)",
        "\u180E": "Mongolian Vowel Separator (U+180E)",
    }

    # ── Layer 3a: Suspicious metadata keywords ───────────────────
    METADATA_KEYWORDS = [
        "ignore", "override", "inject", "instruction", "prompt",
        "execute", "fetch", "disregard", "forget", "pretend",
        "exfil", "http://", "https://", "system", "jailbreak",
    ]

    # ── Layer 4: L33tspeak normalization map ─────────────────────
    L33T_MAP: Dict[str, str] = {
        "0": "o", "1": "i", "3": "e", "4": "a", "5": "s",
        "7": "t", "@": "a", "$": "s", "!": "i", "+": "t",
        "8": "b", "6": "g", "9": "g",
    }

    # ── Scoring ──────────────────────────────────────────────────
    SEVERITY_WEIGHTS = {
        Severity.LOW:      5,
        Severity.MEDIUM:  15,
        Severity.HIGH:    30,
        Severity.CRITICAL:50,
    }

    DEFAULT_BLOCK_THRESHOLD = 30

    # ────────────────────────────────────────────────────────────

    def __init__(self, block_threshold: int = DEFAULT_BLOCK_THRESHOLD):
        self.block_threshold = block_threshold

        # Pre-compile regex patterns
        self._re_injection = [
            (re.compile(pat, re.IGNORECASE | re.DOTALL), desc, sev)
            for pat, desc, sev in self.INJECTION_PATTERNS
        ]
        self._re_exfil = [
            (re.compile(pat, re.IGNORECASE), desc, sev)
            for pat, desc, sev in self.EXFIL_PATTERNS
        ]

    # ── Public API ───────────────────────────────────────────────

    def scan(self, file_path: str) -> ScanResult:
        """
        Auto-detect format and scan the document.

        Supported formats: .pdf, .docx, .doc, .xlsx, .xls, .txt, .md, .csv, .html, .htm
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()
        dispatch = {
            ".pdf":  self._scan_pdf,
            ".docx": self._scan_docx,
            ".doc":  self._scan_docx,
            ".xlsx": self._scan_xlsx,
            ".xls":  self._scan_xlsx,
            ".html": self._scan_html,
            ".htm":  self._scan_html,
            ".txt":  self._scan_plaintext,
            ".md":   self._scan_plaintext,
            ".csv":  self._scan_plaintext,
        }
        handler = dispatch.get(ext)
        if handler is None:
            raise ValueError(f"Unsupported file format: {ext}")
        return handler(file_path)

    def scan_text(self, raw_text: str) -> ScanResult:
        """
        Scan a raw text string (used by the Attack Playground).
        No file I/O — just runs layers 2, 3, and 4 on the text.
        """
        t_start = time.perf_counter()
        findings: List[Finding] = []

        page_map = {1: raw_text}
        findings.extend(self._scan_injection_patterns(raw_text, page_map))
        findings.extend(self._scan_exfil_patterns(raw_text, page_map))
        findings.extend(self._scan_invisible_chars(raw_text))
        findings.extend(self._scan_adversarial(raw_text))

        risk_score = self._calculate_risk(findings)
        verdict    = "BLOCKED" if risk_score >= self.block_threshold else "ALLOWED"
        elapsed_ms = (time.perf_counter() - t_start) * 1000

        return ScanResult(
            filename     = "<playground>",
            file_format  = "text",
            verdict      = verdict,
            risk_score   = min(risk_score, 100),
            scan_time_ms = elapsed_ms,
            findings     = findings,
            total_pages  = 1,
            text_length  = len(raw_text),
        )

    # ── Format-specific scanners ─────────────────────────────────

    def _scan_pdf(self, path: str) -> ScanResult:
        """Scan a PDF file (original 3-layer + new Layer 4)."""
        if not _HAS_PDF:
            raise ImportError("PyPDF2 required: pip install PyPDF2>=3.0.0")

        t_start = time.perf_counter()
        path_obj = Path(path)
        findings: List[Finding] = []

        reader      = PdfReader(str(path_obj))
        total_pages = len(reader.pages)

        # Layer 1 — text extraction
        full_text   = ""
        page_texts: Dict[int, str] = {}
        for idx, page in enumerate(reader.pages):
            extracted = page.extract_text() or ""
            page_texts[idx + 1] = extracted
            full_text += extracted + "\n"

        # Layers 2–4
        findings.extend(self._scan_injection_patterns(full_text, page_texts))
        findings.extend(self._scan_exfil_patterns(full_text, page_texts))
        findings.extend(self._scan_invisible_chars(full_text))
        findings.extend(self._scan_metadata(reader))
        findings.extend(self._scan_text_anomalies(page_texts))
        findings.extend(self._scan_adversarial(full_text))

        return self._build_result(
            path_obj, findings, total_pages, len(full_text),
            t_start, file_format="pdf", metadata_scanned=True,
        )

    def _scan_docx(self, path: str) -> ScanResult:
        """Scan a DOCX file — body paragraphs, tables, and core properties."""
        if not _HAS_DOCX:
            raise ImportError("python-docx required: pip install python-docx>=1.1.0")

        t_start  = time.perf_counter()
        path_obj = Path(path)
        findings: List[Finding] = []

        document = docx.Document(str(path_obj))

        # Extract all text
        paragraphs = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)

        full_text  = "\n".join(paragraphs)
        page_map   = {1: full_text}

        findings.extend(self._scan_injection_patterns(full_text, page_map))
        findings.extend(self._scan_exfil_patterns(full_text, page_map))
        findings.extend(self._scan_invisible_chars(full_text))
        findings.extend(self._scan_text_anomalies(page_map))
        findings.extend(self._scan_adversarial(full_text))

        # Scan core document properties (metadata)
        props = document.core_properties
        meta_text = " ".join(filter(None, [
            props.title, props.author, props.subject, props.description,
            props.keywords, props.category,
        ]))
        if meta_text:
            for _, regex, desc, sev in self._iter_injection(meta_text):
                findings.append(Finding(
                    category="Metadata Injection",
                    description=f"Injection in DOCX property: {desc}",
                    severity=sev,
                    evidence=meta_text[:150],
                ))

        return self._build_result(
            path_obj, findings, 0, len(full_text), t_start, file_format="docx",
        )

    def _scan_xlsx(self, path: str) -> ScanResult:
        """Scan an XLSX file — cell values and formulas."""
        if not _HAS_XLSX:
            raise ImportError("openpyxl required: pip install openpyxl>=3.1.0")

        t_start  = time.perf_counter()
        path_obj = Path(path)
        findings: List[Finding] = []
        texts    = []

        workbook = openpyxl.load_workbook(str(path_obj), data_only=False)
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    val = str(cell.value or "")
                    if val:
                        texts.append(val)
                        # Flag cells with formula that look like injections
                        if val.startswith("=") and len(val) > 5:
                            findings.append(Finding(
                                category="Formula Injection",
                                description="Spreadsheet formula — potential DDE/macro injection vector",
                                severity=Severity.MEDIUM,
                                evidence=val[:100],
                                confidence=0.7,
                            ))

        full_text = "\n".join(texts)
        page_map  = {1: full_text}

        findings.extend(self._scan_injection_patterns(full_text, page_map))
        findings.extend(self._scan_exfil_patterns(full_text, page_map))
        findings.extend(self._scan_invisible_chars(full_text))
        findings.extend(self._scan_adversarial(full_text))

        return self._build_result(
            path_obj, findings, workbook.sheetnames.__len__(),
            len(full_text), t_start, file_format="xlsx",
        )

    def _scan_html(self, path: str) -> ScanResult:
        """Scan an HTML file — visible text, hidden divs, script tags, and meta tags."""
        if not _HAS_BS4:
            raise ImportError("beautifulsoup4 required: pip install beautifulsoup4>=4.12.0")

        t_start  = time.perf_counter()
        path_obj = Path(path)
        findings: List[Finding] = []

        raw_html = path_obj.read_text(encoding="utf-8", errors="replace")
        soup     = BeautifulSoup(raw_html, "html.parser")

        # Check for hidden elements
        for tag in soup.find_all(style=re.compile(r"display\s*:\s*none|visibility\s*:\s*hidden|font-size\s*:\s*0", re.I)):
            hidden_text = tag.get_text(strip=True)
            if hidden_text:
                findings.append(Finding(
                    category="Hidden HTML Content",
                    description="Text hidden via CSS — possible invisible injection payload",
                    severity=Severity.HIGH,
                    evidence=hidden_text[:100],
                    confidence=0.85,
                ))

        # Check script tags
        for script in soup.find_all("script"):
            script_text = script.get_text()
            page_map = {1: script_text}
            findings.extend(self._scan_injection_patterns(script_text, page_map))
            findings.extend(self._scan_exfil_patterns(script_text, page_map))

        # Get all visible text
        full_text = soup.get_text(separator="\n")
        page_map  = {1: full_text}
        findings.extend(self._scan_injection_patterns(full_text, page_map))
        findings.extend(self._scan_exfil_patterns(full_text, page_map))
        findings.extend(self._scan_invisible_chars(full_text))
        findings.extend(self._scan_adversarial(full_text))

        return self._build_result(
            path_obj, findings, 1, len(full_text), t_start, file_format="html",
        )

    def _scan_plaintext(self, path: str) -> ScanResult:
        """Scan a plain text / Markdown / CSV file."""
        t_start  = time.perf_counter()
        path_obj = Path(path)
        findings: List[Finding] = []

        full_text = path_obj.read_text(encoding="utf-8", errors="replace")
        page_map  = {1: full_text}

        findings.extend(self._scan_injection_patterns(full_text, page_map))
        findings.extend(self._scan_exfil_patterns(full_text, page_map))
        findings.extend(self._scan_invisible_chars(full_text))
        findings.extend(self._scan_text_anomalies(page_map))
        findings.extend(self._scan_adversarial(full_text))

        fmt = path_obj.suffix.lower().lstrip(".")
        return self._build_result(
            path_obj, findings, 1, len(full_text), t_start, file_format=fmt,
        )

    # ── Detection layer implementations ─────────────────────────

    def _scan_injection_patterns(
        self, text: str, page_texts: Dict[int, str]
    ) -> List[Finding]:
        """Layer 2a — regex-based prompt injection detection."""
        results = []
        for regex, description, severity in self._re_injection:
            for match in regex.finditer(text):
                page = self._locate_page(match.start(), page_texts)
                results.append(Finding(
                    category   = "Prompt Injection",
                    description= description,
                    severity   = severity,
                    evidence   = match.group(0),
                    page       = page,
                    confidence = 0.95,
                ))
        return results

    def _scan_exfil_patterns(
        self, text: str, page_texts: Dict[int, str]
    ) -> List[Finding]:
        """Layer 2b — data exfiltration URL detection."""
        results = []
        for regex, description, severity in self._re_exfil:
            for match in regex.finditer(text):
                page = self._locate_page(match.start(), page_texts)
                results.append(Finding(
                    category   = "Data Exfiltration",
                    description= description,
                    severity   = severity,
                    evidence   = match.group(0),
                    page       = page,
                    confidence = 0.90,
                ))
        return results

    def _scan_invisible_chars(self, text: str) -> List[Finding]:
        """Layer 2c — detect invisible / zero-width Unicode characters."""
        results  = []
        detected = {}

        for char, name in self.INVISIBLE_CHARS.items():
            count = text.count(char)
            if count > 0:
                detected[name] = count

        if detected:
            total    = sum(detected.values())
            detail   = "; ".join(f"{n}: {c}" for n, c in detected.items())
            severity = Severity.MEDIUM
            if total > 10:
                severity = Severity.HIGH
            if total > 50:
                severity = Severity.CRITICAL

            results.append(Finding(
                category   = "Hidden Characters",
                description= f"Detected {total} invisible Unicode character(s) — "
                             f"possible steganographic payload",
                severity   = severity,
                evidence   = detail,
                confidence = 0.80,
            ))
        return results

    def _scan_metadata(self, reader) -> List[Finding]:
        """Layer 3a — inspect PDF metadata fields for injection."""
        results  = []
        meta     = reader.metadata
        if not meta:
            return results

        field_map = {
            "Title":    getattr(meta, "title", None),
            "Author":   getattr(meta, "author", None),
            "Subject":  getattr(meta, "subject", None),
            "Creator":  getattr(meta, "creator", None),
            "Producer": getattr(meta, "producer", None),
        }

        for field_name, value in field_map.items():
            if not value:
                continue
            value_lower = value.lower()

            for regex, desc, sev in self._re_injection:
                if regex.search(value):
                    results.append(Finding(
                        category   = "Metadata Injection",
                        description= f"Injection payload in PDF '{field_name}' field",
                        severity   = sev,
                        evidence   = f"{field_name}: {value[:150]}",
                        confidence = 0.90,
                    ))
                    break
            else:
                for keyword in self.METADATA_KEYWORDS:
                    if keyword in value_lower:
                        results.append(Finding(
                            category   = "Suspicious Metadata",
                            description= f"Keyword '{keyword}' in PDF '{field_name}' field",
                            severity   = Severity.LOW,
                            evidence   = f"{field_name}: {value[:150]}",
                            confidence = 0.60,
                        ))
                        break
        return results

    def _scan_text_anomalies(self, page_texts: Dict[int, str]) -> List[Finding]:
        """Layer 3b — detect abnormal text patterns."""
        results = []
        for page_num, text in page_texts.items():
            if not text:
                continue
            for line in text.split("\n"):
                stripped = line.strip()
                # Extremely long un-punctuated line — hallmark of encoded payloads
                if len(stripped) > 500 and not any(c in stripped for c in ".!?,;:"):
                    results.append(Finding(
                        category   = "Text Anomaly",
                        description= "Unusually long unpunctuated text — potential hidden payload",
                        severity   = Severity.MEDIUM,
                        evidence   = stripped[:120] + "...",
                        page       = page_num,
                        confidence = 0.70,
                    ))
        return results

    def _scan_adversarial(self, text: str) -> List[Finding]:
        """
        Layer 4 — Adversarial bypass detection:
         (a) Base64-encoded payloads
         (b) L33tspeak normalization
         (c) Unicode lookalike normalization
        """
        findings = []

        # 4a: Base64 decoding
        # Look for long base64-ish strings and try to decode them
        b64_pattern = re.compile(r"[A-Za-z0-9+/]{40,}={0,2}")
        for match in b64_pattern.finditer(text):
            try:
                decoded = base64.b64decode(match.group(0)).decode("utf-8", errors="ignore")
                if len(decoded) > 10:
                    page_map = {1: decoded}
                    sub_findings = self._scan_injection_patterns(decoded, page_map)
                    sub_findings += self._scan_exfil_patterns(decoded, page_map)
                    for f in sub_findings:
                        f.description = "[Base64-decoded] " + f.description
                        f.confidence  = max(0.5, f.confidence - 0.2)
                        f.severity    = Severity(max(1, f.severity.value - 1))  # downgrade one level
                    findings.extend(sub_findings)
            except Exception:
                pass

        # 4b: L33tspeak normalization
        l33t_normalized = text.lower()
        for char, replacement in self.L33T_MAP.items():
            l33t_normalized = l33t_normalized.replace(char, replacement)

        if l33t_normalized != text.lower():
            page_map = {1: l33t_normalized}
            sub = self._scan_injection_patterns(l33t_normalized, page_map)
            for f in sub:
                f.description = "[L33t-decoded] " + f.description
                f.confidence  = max(0.4, f.confidence - 0.3)
            findings.extend(sub)

        # 4c: Unicode NFKD lookalike normalization
        try:
            normalized = unicodedata.normalize("NFKD", text)
            if normalized != text:
                page_map = {1: normalized}
                sub = self._scan_injection_patterns(normalized, page_map)
                for f in sub:
                    f.description = "[Unicode-normalized] " + f.description
                    f.confidence  = max(0.4, f.confidence - 0.25)
                findings.extend(sub)
        except Exception:
            pass

        return findings

    # ── Utilities ────────────────────────────────────────────────

    def _iter_injection(self, text: str):
        """Yield (match, regex, desc, sev) for injection patterns in text."""
        for regex, desc, sev in self._re_injection:
            for match in regex.finditer(text):
                yield match, regex, desc, sev

    def _build_result(
        self,
        path_obj:        Path,
        findings:        List[Finding],
        total_pages:     int,
        text_length:     int,
        t_start:         float,
        file_format:     str = "unknown",
        metadata_scanned:bool = False,
    ) -> ScanResult:
        risk_score = self._calculate_risk(findings)
        verdict    = "BLOCKED" if risk_score >= self.block_threshold else "ALLOWED"
        elapsed_ms = (time.perf_counter() - t_start) * 1000

        return ScanResult(
            filename         = path_obj.name,
            file_format      = file_format,
            verdict          = verdict,
            risk_score       = min(risk_score, 100),
            scan_time_ms     = elapsed_ms,
            findings         = findings,
            total_pages      = total_pages,
            text_length      = text_length,
            metadata_scanned = metadata_scanned,
        )

    def _locate_page(
        self, char_offset: int, page_texts: Dict[int, str]
    ) -> Optional[int]:
        """Map a char offset in concatenated text back to a page number."""
        running = 0
        for page_num in sorted(page_texts.keys()):
            page_len = len(page_texts[page_num]) + 1
            if running + page_len > char_offset:
                return page_num
            running += page_len
        return None

    def _calculate_risk(self, findings: List[Finding]) -> int:
        """Compute a composite risk score (0–100) weighted by severity and confidence."""
        if not findings:
            return 0
        raw = sum(
            self.SEVERITY_WEIGHTS.get(f.severity, 0) * f.confidence
            for f in findings
        )
        return int(min(raw, 100))
